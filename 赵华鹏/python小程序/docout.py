import os
import json
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm,Pt
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment,Side,Border

def outXlsx(fDict,name):
    #fDict:JSON文件中的字典
    #name:字典中的键
    #函数作用：输出.xlsx文件
    m_dict = fDict[name]
    workbook = Workbook()
    sheet_names = workbook.sheetnames
    print(sheet_names)
    sheet = workbook.active
    sheet.title = name
    list1 = ['技术能手（' + name + '）管理月报（一）']
    sheet.append(list1)
    sheet.merge_cells("A1:F1")

    my_dict = m_dict["技术成果"]
    list1 = ['技术成果','技术成果','季度目标',my_dict['季度目标']]
    sheet.append(list1)
    sheet.merge_cells("D2:F2")
    sheet.row_dimensions[2].height = 40

    list1 = ['技术成果','技术成果','月度目标',my_dict['月度目标'],'自我评分','复核得分']
    sheet.append(list1)

    list1 = ['技术成果','技术成果','完成情况',my_dict['完成情况']]
    sheet.append(list1)
    sheet.merge_cells("A2:B4")

    my_dict = m_dict["重要活动"]
    list1 = ['重要活动','重要活动','参加时间','活动内容','负责人签确','得分']
    sheet.append(list1)

    for key,value in my_dict.items():
        list1 = ['重要活动','重要活动']
        list1.append(key)
        list1.append(value)
        sheet.append(list1)

    list1 = ['重要活动','重要活动','得分小计']
    sheet.append(list1)

    sheet.column_dimensions["A"].width = 3
    sheet.column_dimensions["B"].width = 3
    sheet.column_dimensions["C"].width = 10
    sheet.column_dimensions["D"].width = 40 
    sheet.column_dimensions["E"].width = 10
    sheet.column_dimensions["F"].width = 10
    sheet.merge_cells("A5:B10")
    sheet.merge_cells("C10:D10")

    list1 = ['技术能手（' + name + '）管理月报（二）']
    sheet.append(list1)
    sheet.merge_cells("A11:F11")

    list1 = ['学习传授','自我学习','时间','内容','内容','得分']
    sheet.append(list1)

    my_dict = m_dict["自我学习"]
    for key,value in my_dict.items():
        list1 = ['学习传授','自我学习']
        list1.append(key)
        list1.append(value)
        sheet.append(list1)
    list1 = ['学习传授','自我学习','得分小计']
    sheet.append(list1)
    sheet.merge_cells("C23:E23")
    sheet.merge_cells("B12:B23")

    list1 = ['学习传授','授课','时间','内容','参与人员','得分']
    sheet.append(list1)
    my_dict = m_dict["授课"]
    for key,value in my_dict.items():
        list1 = ['学习传授','授课']
        list1.append(key)
        list1.append(value)
        sheet.append(list1)
    list1 = ['学习传授','授课','得分小计']
    sheet.append(list1)
    sheet.merge_cells("C27:E27")
    sheet.merge_cells("B24:B27")

    list1 = ['学习传授','师带徒','时间','内容','徒弟签确','得分']
    sheet.append(list1)
    my_dict = m_dict["师带徒"]
    for key,value in my_dict.items():
        list1 = ['学习传授','师带徒']
        list1.append(key)
        list1.append(value)
        sheet.append(list1)
    list1 = ['学习传授','师带徒','得分小计']
    sheet.append(list1)
    sheet.merge_cells("C39:E39")
    sheet.merge_cells("B28:B39")
    sheet.merge_cells("A12:A39")

    list1 = ['月度合计得分']
    sheet.append(list1)
    sheet.merge_cells("A40:E40")

    cell = sheet["A1:F40"]
    alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
    side1 = Side(style="thin",color="000000")
    border = Border(left=side1,right=side1,top=side1,bottom=side1)
    for i in cell:
        for j in i:
            j.alignment = alignment
            j.border = border

    workbook.save(filename = r"2025.xlsx")

def outDocx(fDict,name):
    #fDict:JSON文件中的字典
    #name:字典中的键
    #函数作用：输出.Docx文件
    m_dict = fDict[name]
    doc = docx.Document()
    doc.styles['Normal'].font.name= u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    paragraph_title = doc.add_paragraph()
    paragraph_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#水平居中
    run = paragraph_title.add_run("技术能手（"+name+"）管理月报（一）")
    run.font.name = '仿宋'
    run.font.size = Pt(18)
    table = doc.add_table(rows=9,cols=5,style='Table Grid')
    table.style.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER#水平居中
    table.autofit = False
    #table.style.font.bold = True
    table.style.font.name = u'仿宋'
    table.style.font.size = Pt(14)
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(2.5)
    table.columns[2].width = Cm(5)
    table.columns[3].width = Cm(2)
    table.rows[0].height = Cm(2)
    table.rows[1].height = Cm(2)
     
    my_dict = m_dict["技术成果"]
    print(my_dict)
    merged_cell = table.cell(0,0).merge(table.cell(2,0))
    merged_cell.text = '技术成果'
    cell =table.cell(0,1)
    cell.text = '季度目标'
    merged_cell = table.cell(0,2).merge(table.cell(0,4))
    merged_cell.text = my_dict['季度目标']
    cell =table.cell(1,1)
    cell.text = '月度目标'
    cell =table.cell(1,2)
    cell.text = my_dict['月度目标']
    cell =table.cell(1,3)
    cell.text = '自我评分'
    cell =table.cell(1,4)
    cell.text = '复核得分'
    cell =table.cell(2,1)
    cell.text = '完成情况'
    cell =table.cell(2,2)
    cell.text = my_dict['完成情况']
    
    my_dict = m_dict["重要活动"]
    print(my_dict)
    merged_cell = table.cell(3,0).merge(table.cell(8,0))
    merged_cell.text = '重要活动'
    merged_cell = table.cell(8,1).merge(table.cell(8,3))
    merged_cell.text = '得分小计'
    cell = table.cell(3,1)
    cell.text = '参加时间'
    cell = table.cell(3,2)
    cell.text = '活动内容'
    cell = table.cell(3,3)
    cell.text = '负责人签确'
    cell = table.cell(3,4)
    cell.text = '得分'
    i = 4
    for key,value in my_dict.items():
        cell = table.cell(i,1)
        cell.text = key
        cell = table.cell(i,2)
        cell.text = value
        i = i + 1

    for r in range(len(table.rows)):#循环将每一行，每一列都设置为居中
        for c in range(len(table.columns)):
            #table.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER#水平居中
            table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER#垂直居中
    #第二页
    
    paragraph_title = doc.add_paragraph()
    paragraph_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#水平居中
    run = paragraph_title.add_run("技术能手（"+name+"）管理月报（二）")
    run.font.name = '仿宋'
    run.font.size = Pt(18)
    table = doc.add_table(rows=29,cols=6,style='Table Grid')
    table.autofit = False
    table.columns[0].width = Cm(1)
    table.columns[1].width = Cm(1)
    table.columns[2].width = Cm(3)
    table.columns[3].width = Cm(5)
    table.columns[4].width = Cm(2)
    
    merged_cell = table.cell(0,0).merge(table.cell(27,0))
    merged_cell.text = '学习传授'
    
    merged_cell = table.cell(0,1).merge(table.cell(11,1))
    merged_cell.text = '自我学习'
    cell = table.cell(0,2)
    cell.text = '时间'
    merged_cell = table.cell(0,3).merge(table.cell(0,4))
    merged_cell.text = '内容'
    cell = table.cell(0,5)
    cell.text = '得分'
    merged_cell = table.cell(11,2).merge(table.cell(11,4))
    merged_cell.text = '得分小计'
    
    my_dict = m_dict["自我学习"]
    print(my_dict)
    
    i = 1
    for key,value in my_dict.items():
        cell = table.cell(i,2)
        cell.text = key
        merged_cell = table.cell(i,3).merge(table.cell(i,4))
        merged_cell.text = value
        i = i + 1
    
    
    merged_cell = table.cell(12,1).merge(table.cell(15,1))
    merged_cell.text = '授课'
    cell = table.cell(12,2)
    cell.text = '时间'
    cell = table.cell(12,3)
    cell.text = '内容'
    cell = table.cell(12,4)
    cell.text = '参与人员'
    cell = table.cell(12,5)
    cell.text = '得分'
    merged_cell = table.cell(15,2).merge(table.cell(15,4))
    merged_cell.text = '得分小计'
    
    my_dict = m_dict["授课"]
    print(my_dict)
    
    i = 13
    for key,value in my_dict.items():
        cell = table.cell(i,2)
        cell.text = key
        cell = table.cell(i,3)
        cell.text = value 
        i = i + 1
    
    merged_cell = table.cell(16,1).merge(table.cell(27,1))
    merged_cell.text = '授课'
    cell = table.cell(16,2)
    cell.text = '时间'
    cell = table.cell(16,3)
    cell.text = '内容'
    cell = table.cell(16,4)
    cell.text = '徒弟签确'
    cell = table.cell(16,5)
    cell.text = '得分'
    merged_cell = table.cell(27,2).merge(table.cell(27,4))
    merged_cell.text = '得分小计'
    merged_cell = table.cell(28,0).merge(table.cell(28,4))
    merged_cell.text = '月度合计得分'
    
    my_dict = m_dict["师带徒"]
    print(my_dict)
    
    i = 17
    for key,value in my_dict.items():
        cell = table.cell(i,2)
        cell.text = key
        cell = table.cell(i,3)
        cell.text = value 
        i = i + 1
    for r in range(len(table.rows)):#循环将每一行，每一列都设置为居中
        for c in range(len(table.columns)):
            table.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER#水平居中
            table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER#垂直居中
    
    
    doc.save(r"out.docx")

with open('f.json','r',encoding='utf-8') as f:
    json_dict = json.load(f)
keys_list = list(json_dict.keys())
for i in range(len(keys_list)):
    print(str(i)+"->"+keys_list[i])
b = int(input("按提示输入一个数字来选择要生成文件:"))
print(keys_list[b])
outDocx(json_dict,keys_list[b])
#outXlsx(json_dict,keys_list[0])
#outXlsx(json_dict,keys_list[1])


